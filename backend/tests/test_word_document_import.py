from docx import Document as WordDocument

from app.services.resource_service import extract_docx_text


def test_extract_docx_text_reads_paragraphs_and_tables(tmp_path):
    path = tmp_path / "course-notes.docx"
    document = WordDocument()
    document.add_paragraph("向量空间的加法满足结合律。")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "概念"
    table.cell(0, 1).text = "说明"
    table.cell(1, 0).text = "线性组合"
    table.cell(1, 1).text = "向量的标量倍数之和"
    document.save(path)

    extracted = extract_docx_text(path)

    assert "向量空间的加法满足结合律。" in extracted
    assert "概念 | 说明" in extracted
    assert "线性组合 | 向量的标量倍数之和" in extracted
