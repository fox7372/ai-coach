from app.runtime import *
from app.schemas import *
from app.services.retrieval import index_chunks_in_chroma, make_embeddings

def clean_subtitle_text(text_value: str) -> str:
    without_tags = re.sub(r"<[^>]+>", "", text_value)
    return re.sub(r"\s+", " ", html.unescape(without_tags)).strip()


def parse_timestamp(value: str) -> float:
    parts = value.replace(",", ".").split(":")
    if len(parts) == 3:
        hours, minutes, seconds = parts
    elif len(parts) == 2:
        hours, minutes, seconds = "0", parts[0], parts[1]
    else:
        return 0.0
    return int(hours) * 3600 + int(minutes) * 60 + float(seconds)


def parse_subtitle_segments(content: str) -> list[dict[str, object]]:
    normalized = content.replace("\ufeff", "").replace("\r\n", "\n").replace("\r", "\n")
    blocks = re.split(r"\n\s*\n", normalized)
    segments: list[dict[str, object]] = []
    time_pattern = re.compile(r"(?P<start>\d{1,2}:\d{2}(?::\d{2})?[\.,]\d{1,3})\s*-->\s*(?P<end>\d{1,2}:\d{2}(?::\d{2})?[\.,]\d{1,3})")
    for block in blocks:
        lines = [line.strip() for line in block.splitlines() if line.strip() and not line.strip().isdigit()]
        time_index = next((idx for idx, line in enumerate(lines) if "-->" in line), -1)
        if time_index < 0:
            continue
        match = time_pattern.search(lines[time_index])
        if not match:
            continue
        text_lines = [line for line in lines[time_index + 1:] if not line.startswith(("NOTE", "STYLE", "WEBVTT"))]
        segment_text = clean_subtitle_text(" ".join(text_lines))
        if segment_text:
            segments.append({"start": parse_timestamp(match.group("start")), "end": parse_timestamp(match.group("end")), "text": segment_text})
    return segments


def merge_segments_to_chunks(segments: list[dict[str, object]], target_chars: int = 800, max_seconds: int = 120) -> list[dict[str, object]]:
    chunks: list[dict[str, object]] = []
    current: list[dict[str, object]] = []
    for segment in segments:
        current.append(segment)
        start = float(current[0]["start"])
        end = float(current[-1]["end"])
        text_value = " ".join(str(item["text"]) for item in current)
        if len(text_value) >= target_chars or end - start >= max_seconds:
            chunks.append({"start": start, "end": end, "text": text_value})
            current = current[-1:]
    if current:
        chunks.append({"start": float(current[0]["start"]), "end": float(current[-1]["end"]), "text": " ".join(str(item["text"]) for item in current)})
    return chunks


def fetch_text_url(url: str) -> str:
    validate_public_url(url)
    request = Request(url, headers={"User-Agent": "AI-Learning-MVP/0.1"})
    with urlopen(request, timeout=20) as response:
        charset = response.headers.get_content_charset() or "utf-8"
        return response.read().decode(charset, errors="replace")


def fetch_webpage_html(url: str) -> str:
    clean_url = validate_public_url(url)
    request = Request(
        clean_url,
        headers={
            "User-Agent": "Mozilla/5.0 (compatible; AI-Learning-MVP/0.1; +https://localhost)",
            "Accept": "text/html,application/xhtml+xml",
            "Accept-Language": "zh-CN,zh;q=0.9,en;q=0.8",
        },
    )
    with urlopen(request, timeout=25) as response:
        content_type = response.headers.get("Content-Type", "")
        if "text/html" not in content_type and "application/xhtml" not in content_type:
            raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="该链接不是 HTML 网页")
        charset = response.headers.get_content_charset() or "utf-8"
        return response.read().decode(charset, errors="replace")


def extract_webpage_content(url: str) -> dict[str, object]:
    html_content = fetch_webpage_html(url)
    title_match = re.search(r"<title[^>]*>(.*?)</title>", html_content, flags=re.IGNORECASE | re.DOTALL)
    title = clean_subtitle_text(title_match.group(1)) if title_match else url

    without_noise = re.sub(r"<(script|style|noscript|svg|canvas)[^>]*>.*?</\1>", " ", html_content, flags=re.IGNORECASE | re.DOTALL)
    without_comments = re.sub(r"<!--.*?-->", " ", without_noise, flags=re.DOTALL)
    text_content = re.sub(r"</(p|div|section|article|li|h[1-6]|tr)>", "\n", without_comments, flags=re.IGNORECASE)
    text_content = re.sub(r"<[^>]+>", " ", text_content)
    text_content = html.unescape(text_content)
    lines = [re.sub(r"\s+", " ", line).strip() for line in text_content.splitlines()]
    lines = [line for line in lines if len(line) >= 2]
    body_text = "\n".join(lines)

    links = []
    parsed_base = urlparse(url)
    for href, label in re.findall(r"<a[^>]+href=[\"']([^\"']+)[\"'][^>]*>(.*?)</a>", html_content, flags=re.IGNORECASE | re.DOTALL):
        label_text = clean_subtitle_text(label)
        if not label_text:
            continue
        if href.startswith("/"):
            href = f"{parsed_base.scheme}://{parsed_base.netloc}{href}"
        elif href.startswith("#"):
            href = f"{url.rstrip('/')}{href}"
        if href.startswith(("http://", "https://")):
            links.append({"title": label_text[:120], "url": href})
        if len(links) >= 30:
            break

    if len(body_text) < 100:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="网页正文太少，可能需要登录或由脚本动态加载")

    return {"title": title[:255], "text": body_text, "links": links}


def split_text_to_chunks(text_value: str, target_chars: int = 1200, overlap_chars: int = 120) -> list[str]:
    paragraphs = [item.strip() for item in re.split(r"\n{1,}", text_value) if item.strip()]
    chunks: list[str] = []
    current = ""
    for paragraph in paragraphs:
        if len(current) + len(paragraph) + 1 > target_chars and current:
            chunks.append(current.strip())
            current = current[-overlap_chars:] if overlap_chars else ""
        current = f"{current}\n{paragraph}".strip()
    if current:
        chunks.append(current.strip())
    return chunks


def normalize_for_page_match(text_value: str) -> str:
    return re.sub(r"\s+", "", text_value).lower()


def extract_pdf_pages(path: Path) -> list[dict[str, object]]:
    try:
        import fitz
    except ImportError as exc:
        raise RuntimeError("后端缺少 PyMuPDF，请先安装 requirements.txt") from exc

    pages: list[dict[str, object]] = []
    with fitz.open(path) as pdf:
        for page_index, page in enumerate(pdf, start=1):
            text_value = page.get_text("text").strip()
            if text_value:
                pages.append({"page_number": page_index, "text": text_value})
    if not pages:
        raise RuntimeError("PDF 没有提取到可用文本，可能是扫描版图片 PDF，需要 OCR。")
    return pages


def extract_document_with_docling(path: Path, input_format: object | None = None) -> str:
    try:
        from docling.datamodel.base_models import InputFormat
        from docling.datamodel.pipeline_options import PdfPipelineOptions
        from docling.document_converter import DocumentConverter
        from docling.document_converter import PdfFormatOption
    except ImportError as exc:
        raise RuntimeError("后端缺少 Docling，请先安装 requirements.txt") from exc

    pipeline_options = PdfPipelineOptions()
    pipeline_options.document_timeout = 90
    pipeline_options.do_ocr = False
    format_options = {}
    if input_format is None or input_format == InputFormat.PDF:
        format_options[InputFormat.PDF] = PdfFormatOption(pipeline_options=pipeline_options)
    converter = DocumentConverter(format_options=format_options)
    result = converter.convert(str(path))
    markdown = result.document.export_to_markdown().strip()
    if not markdown:
        raise RuntimeError("Docling 没有从资料中提取到可用文本。")
    return markdown


def extract_pdf_with_docling(path: Path) -> str:
    return extract_document_with_docling(path)


def find_chunk_page_number(chunk_text: str, pages: list[dict[str, object]]) -> int | None:
    normalized_chunk = normalize_for_page_match(chunk_text)
    if not normalized_chunk:
        return None

    probe = normalized_chunk[: min(len(normalized_chunk), 80)]
    if len(probe) < 20:
        return None

    for page in pages:
        page_text = normalize_for_page_match(str(page["text"]))
        if probe in page_text:
            return int(page["page_number"])
    return None


def save_document_chunks(
    db: Session,
    document: Document,
    structured_text: str,
    parser_name: str,
    source_type: str,
    pages: list[dict[str, object]] | None = None,
) -> int:
    try:
        rag_service.delete_document(document.id)
    except Exception:
        pass
    db.execute(delete(DocumentChunk).where(DocumentChunk.document_id == document.id))
    chunk_index = 0
    chunks = split_text_to_chunks(structured_text)
    if not chunks:
        raise RuntimeError("资料没有切分出可用片段。")

    embeddings = make_embeddings(chunks)
    pending_index: list[DocumentChunk] = []
    for chunk_text, embedding in zip(chunks, embeddings):
        page_number = find_chunk_page_number(chunk_text, pages or []) if pages else None
        if page_number:
            section_title = f"第 {page_number} 页片段"
        elif source_type == "presentation":
            section_title = f"Docling 演示文稿片段 {chunk_index + 1}"
        elif parser_name == "pymupdf":
            section_title = f"PyMuPDF 片段 {chunk_index + 1}"
        elif parser_name == "docling":
            section_title = f"Docling 结构化片段 {chunk_index + 1}"
        else:
            section_title = f"资料片段 {chunk_index + 1}"
        metadata = {
            "resource_title": document.filename,
            "source_type": source_type,
            "parser": parser_name,
        }
        if page_number:
            metadata["page_number"] = page_number

        chunk = DocumentChunk(
            document_id=document.id,
            course_id=document.course_id,
            chunk_text=chunk_text,
            page_number=page_number,
            start_time=None,
            end_time=None,
            section_title=section_title,
            source_url=None,
            metadata_json=json.dumps(metadata, ensure_ascii=False),
            token_count=len(chunk_text),
            chunk_index=chunk_index,
            embedding=json.dumps(embedding, ensure_ascii=False),
        )
        db.add(chunk)
        db.flush()
        pending_index.append(chunk)
        chunk_index += 1
    index_chunks_in_chroma(pending_index)
    return chunk_index


def process_pdf_document(db: Session, document: Document, parser: str = "pymupdf") -> None:
    document_id = document.id
    document.status = "processing"
    document.parse_status = "processing"
    document.progress_stage = f"{parser}_extract_pdf"
    document.error_message = None
    db.commit()

    try:
        path = Path(document.storage_path)
        pages: list[dict[str, object]] = []
        parser_name = parser if parser in {"docling", "pymupdf"} else "pymupdf"
        if parser_name == "pymupdf":
            pages = extract_pdf_pages(path)
            structured_text = "\n\n".join(str(page["text"]) for page in pages)
        else:
            try:
                structured_text = extract_pdf_with_docling(path)
            except Exception:
                pages = extract_pdf_pages(path)
                structured_text = "\n\n".join(str(page["text"]) for page in pages)
                parser_name = "pymupdf"
            else:
                try:
                    pages = extract_pdf_pages(path)
                except Exception:
                    pages = []

        chunk_index = save_document_chunks(db, document, structured_text, parser_name, "pdf", pages)

        document.raw_content = structured_text
        document.status = "ready"
        document.parse_status = "ready"
        document.progress_stage = f"completed:{chunk_index} chunks:{parser_name}"
        db.commit()
    except Exception as exc:
        mark_document_failed(db, document_id, exc)


def process_presentation_document(db: Session, document: Document) -> None:
    document_id = document.id
    document.status = "processing"
    document.parse_status = "processing"
    document.progress_stage = "docling_extract_presentation"
    document.error_message = None
    db.commit()

    try:
        structured_text = extract_document_with_docling(Path(document.storage_path))
        chunk_count = save_document_chunks(db, document, structured_text, "docling", "presentation")
        document.raw_content = structured_text
        document.status = "ready"
        document.parse_status = "ready"
        document.progress_stage = f"completed:{chunk_count} chunks:docling"
        db.commit()
    except Exception as exc:
        mark_document_failed(db, document_id, exc)


def extract_docx_text(path: Path) -> str:
    try:
        from docx import Document as WordDocument
    except ImportError as exc:
        raise RuntimeError("后端缺少 python-docx，请先安装 requirements.txt") from exc

    word_document = WordDocument(path)
    parts = [paragraph.text.strip() for paragraph in word_document.paragraphs if paragraph.text.strip()]
    for table in word_document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    text_value = "\n\n".join(parts).strip()
    if not text_value:
        raise RuntimeError("Word 文档没有提取到可用文本。")
    return text_value


def process_word_document(db: Session, document: Document) -> None:
    document_id = document.id
    document.status = "processing"
    document.parse_status = "processing"
    document.progress_stage = "python_docx_extract"
    document.error_message = None
    db.commit()

    try:
        structured_text = extract_docx_text(Path(document.storage_path))
        chunk_count = save_document_chunks(db, document, structured_text, "python-docx", "word")
        document.raw_content = structured_text
        document.status = "ready"
        document.parse_status = "ready"
        document.progress_stage = f"completed:{chunk_count} chunks:python-docx"
        db.commit()
    except Exception as exc:
        mark_document_failed(db, document_id, exc)


def mark_document_failed(db: Session, document_id: int, exc: Exception) -> None:
    """Restore a usable transaction before persisting document failure details."""
    db.rollback()
    document = db.get(Document, document_id)
    if document is None:
        return
    document.status = "failed"
    document.parse_status = "failed"
    document.progress_stage = "failed"
    document.error_message = clean_error_message(exc)
    db.commit()


def choose_subtitle_track(info: dict[str, object], preferred_language: str) -> tuple[str | None, str | None, str | None]:
    language_candidates = [preferred_language, "zh-CN", "zh-Hans", "zh", "en"]
    for subtitle_type, key in (("manual", "subtitles"), ("automatic", "automatic_captions")):
        tracks = info.get(key)
        if not isinstance(tracks, dict):
            continue
        for language in language_candidates:
            variants = tracks.get(language)
            if not isinstance(variants, list):
                continue
            selected = next((item for item in variants if item.get("ext") in {"vtt", "srt"} and item.get("url")), None)
            if selected:
                return str(selected["url"]), language, subtitle_type
    return None, None, None


def document_to_resource(db: Session, document: Document) -> ResourceOut:
    chunk_count = db.scalar(select(func.count()).select_from(DocumentChunk).where(DocumentChunk.document_id == document.id)) or 0
    return ResourceOut(
        id=document.id,
        course_id=document.course_id,
        title=document.filename,
        source_type=document.source_type or document.file_type,
        source_url=document.source_url,
        platform=document.platform,
        author=document.author,
        duration_seconds=document.duration_seconds,
        thumbnail_url=document.thumbnail_url,
        language=document.language,
        subtitle_type=document.subtitle_type,
        status=document.status or document.parse_status,
        progress_stage=document.progress_stage,
        priority=document.priority,
        error_message=document.error_message,
        chunk_count=int(chunk_count),
    )


def clean_error_message(exc: Exception) -> str:
    return re.sub(r"\x1b\[[0-9;]*m", "", str(exc)).strip()


def extract_mistake_question_text(text_value: str | None) -> str:
    text_value = repair_mojibake(text_value or "")
    for raw_line in text_value.splitlines():
        line = raw_line.strip(" -*\t")
        plain_line = re.sub(r"[*_`#]+", "", line).strip()
        if "题目" in plain_line and ("：" in plain_line or ":" in plain_line):
            separator = "：" if "：" in plain_line else ":"
            question_text = plain_line.split(separator, 1)[1].strip()
            if question_text:
                return question_text
    return ""
